from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "resume"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "MD_Mirza_Galib_Palash_Resume.docx"
PDF_PATH = OUT_DIR / "MD_Mirza_Galib_Palash_Resume.pdf"
PUBLIC_PDF_PATH = ROOT / "public" / "resume.pdf"

ACCENT = RGBColor(31, 78, 121)
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(92, 92, 92)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, size=9.5, bold=False, color=INK):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def section_title(doc, title):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=7, after=2)
    r = add_run(p, title.upper(), size=10.5, bold=True, color=ACCENT)
    r.font.all_caps = True
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    border.append(bottom)
    p_pr.append(border)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph_spacing(p, before=0, after=1.5, line=1.0)
    add_run(p, text, size=8.8, color=INK)
    return p


def role_line(doc, title, meta):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=0)
    add_run(p, title, size=9.3, bold=True, color=INK)
    add_run(p, f" | {meta}", size=8.6, color=MUTED)
    return p


def make_resume():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.38)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=0)
    add_run(title, "MD MIRZA GALIB PALASH", size=18, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=1)
    add_run(subtitle, "Full Stack Developer", size=10.5, bold=True, color=INK)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, after=4)
    add_run(
        contact,
        "Dhaka, Mohammadpur | 01577088342 | mirza.galib.palash@gmail.com | mirzagalib.xyz",
        size=8.7,
        color=MUTED,
    )
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(links, after=5)
    add_run(
        links,
        "LinkedIn: www.linkedin.com/in/md-mirza-galib-palash | GitHub: github.com/GalibDev",
        size=8.5,
        color=MUTED,
    )

    section_title(doc, "Professional Summary")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2, line=1.05)
    add_run(
        p,
        "Motivated full stack developer focused on building responsive, scalable, and user-friendly web applications. "
        "Experienced with React, Next.js, JavaScript, Tailwind CSS, Supabase, MongoDB, and Node.js through real portfolio, e-commerce, and client-style projects. "
        "Comfortable turning ideas into clean interfaces, practical features, and deployable products.",
        size=8.9,
        color=INK,
    )

    section_title(doc, "Technical Skills")
    skills = [
        ("Frontend", "React.js, Next.js, JavaScript, TypeScript, Tailwind CSS, HTML5, CSS3"),
        ("Backend", "Node.js, Express.js, REST APIs, Supabase"),
        ("Database", "MongoDB, PostgreSQL/Supabase"),
        ("Tools", "Git, GitHub, Vercel, VS Code, responsive design, deployment"),
    ]
    table = doc.add_table(rows=len(skills), cols=2)
    table.autofit = False
    for row, (label, value) in zip(table.rows, skills):
      row.cells[0].width = Inches(1.15)
      row.cells[1].width = Inches(6.2)
      for cell in row.cells:
          cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
          remove_cell_border(cell)
          for para in cell.paragraphs:
              set_paragraph_spacing(para, after=0)
      add_run(row.cells[0].paragraphs[0], label, size=8.8, bold=True, color=ACCENT)
      add_run(row.cells[1].paragraphs[0], value, size=8.8, color=INK)

    section_title(doc, "Projects")
    role_line(doc, "Jersey Shop E-commerce Website", "Next.js, Tailwind CSS, Supabase | Live: nivalo.xyz")
    bullet(doc, "Built a responsive jersey e-commerce experience with product presentation, modern UI sections, and clean shopping flow.")
    bullet(doc, "Focused on mobile-first layout, attractive product cards, and production deployment.")

    role_line(
        doc,
        "IdeaVault - Idea Management Web App",
        "React/Next.js, Node.js, MongoDB | Live: idea-vault-client-sigma.vercel.app",
    )
    bullet(doc, "Developed a secure idea management platform where users can save, organize, and manage project ideas.")
    bullet(
        doc,
        "Client repo: github.com/GalibDev/idea-vault-client | Server repo: github.com/GalibDev/idea-vault-server",
    )

    role_line(doc, "Gatrix Robotics Website", "Modern responsive website | Live: gatrix.xyz")
    bullet(doc, "Created a polished robotics/business website with clear content sections, responsive layout, and deployment-ready UI.")

    section_title(doc, "Experience")
    role_line(doc, "Freelance / Local Client Web Projects", "Full Stack / Frontend Developer | 2024-Present")
    bullet(doc, "Designed and delivered 2-3 small websites for local clients, including responsive layouts and deployment support.")
    bullet(doc, "Worked directly with clients to understand requirements, refine UI, and publish finished websites.")

    section_title(doc, "Education")
    edu = [
        ("BSc in CSE", "Shyamoli Engineering College, Affiliated with University of Dhaka | 2024-Present"),
        ("HSC - Science", "Rangpur Model College, Rangpur | 2023-2024 | GPA 5.00"),
        ("SSC - Science", "Shishu Niketon High School | 2021-2022 | GPA 5.00"),
    ]
    for degree, meta in edu:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=1, after=0)
        add_run(p, degree, size=8.9, bold=True, color=INK)
        add_run(p, f" | {meta}", size=8.6, color=MUTED)

    section_title(doc, "Additional")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0)
    add_run(
        p,
        "Languages: Bangla, English | Interests: full stack development, SaaS products, e-commerce, and robotics websites.",
        size=8.8,
        color=INK,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def pstyle(name, size=9, leading=None, color=colors.HexColor("#1f1f1f"), bold=False, align=0):
    return ParagraphStyle(
        name=name,
        fontName="Helvetica-Bold" if bold else "Helvetica",
        fontSize=size,
        leading=leading or size + 1.4,
        textColor=color,
        alignment=align,
        spaceAfter=0,
        spaceBefore=0,
    )


def project_block(title, meta, bullets):
    styles = make_pdf_styles()
    block = [
        Paragraph(f"<b>{title}</b> <font color='#666666'>| {meta}</font>", styles["role"]),
        Spacer(1, 1.5),
        ListFlowable(
            [ListItem(Paragraph(item, styles["bullet"]), leftIndent=8) for item in bullets],
            bulletType="bullet",
            start="circle",
            leftIndent=12,
            bulletFontSize=5.5,
            bulletOffsetY=1,
        ),
    ]
    return block


def make_pdf_styles():
    return {
        "name": pstyle("Name", size=17, leading=19, color=colors.HexColor("#1f4e79"), bold=True, align=1),
        "subtitle": pstyle("Subtitle", size=10.5, leading=12, bold=True, align=1),
        "contact": pstyle("Contact", size=8.4, leading=10, color=colors.HexColor("#5c5c5c"), align=1),
        "body": pstyle("Body", size=8.65, leading=10.2),
        "section": pstyle("Section", size=10, leading=11.5, color=colors.HexColor("#1f4e79"), bold=True),
        "role": pstyle("Role", size=8.8, leading=10.2),
        "bullet": pstyle("Bullet", size=8.3, leading=9.6),
        "muted": pstyle("Muted", size=8.2, leading=9.4, color=colors.HexColor("#666666")),
    }


def add_section(story, title, styles):
    story.append(Spacer(1, 5))
    story.append(Paragraph(title.upper(), styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#d9e2f3"), spaceBefore=1, spaceAfter=3))


def make_pdf_resume():
    styles = make_pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.46 * inch,
        rightMargin=0.46 * inch,
        topMargin=0.38 * inch,
        bottomMargin=0.35 * inch,
    )

    story = [
        Paragraph("MD MIRZA GALIB PALASH", styles["name"]),
        Paragraph("Full Stack Developer", styles["subtitle"]),
        Spacer(1, 2),
        Paragraph("Dhaka, Mohammadpur | 01577088342 | mirza.galib.palash@gmail.com | mirzagalib.xyz", styles["contact"]),
        Paragraph("LinkedIn: www.linkedin.com/in/md-mirza-galib-palash | GitHub: github.com/GalibDev", styles["contact"]),
    ]

    add_section(story, "Professional Summary", styles)
    story.append(
        Paragraph(
            "Motivated full stack developer focused on building responsive, scalable, and user-friendly web applications. "
            "Experienced with React, Next.js, JavaScript, TypeScript, Tailwind CSS, Supabase, MongoDB, and Node.js through real portfolio, e-commerce, and client-style projects. "
            "Comfortable turning ideas into clean interfaces, practical features, and deployable products.",
            styles["body"],
        )
    )

    add_section(story, "Technical Skills", styles)
    skills = [
        ["Frontend", "React.js, Next.js, JavaScript, TypeScript, Tailwind CSS, HTML5, CSS3"],
        ["Backend", "Node.js, Express.js, REST APIs, Supabase"],
        ["Database", "MongoDB, PostgreSQL/Supabase"],
        ["Tools", "Git, GitHub, Vercel, VS Code, responsive design, deployment"],
    ]
    table = Table(
        [[Paragraph(f"<b>{label}</b>", styles["body"]), Paragraph(value, styles["body"])] for label, value in skills],
        colWidths=[1.05 * inch, 6.35 * inch],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1f4e79")),
            ]
        )
    )
    story.append(table)

    add_section(story, "Projects", styles)
    for item in project_block(
        "Jersey Shop E-commerce Website",
        "Next.js, Tailwind CSS, Supabase | Live: nivalo.xyz",
        [
            "Built a responsive jersey e-commerce experience with product presentation, modern UI sections, and a clean shopping flow.",
            "Focused on mobile-first layout, attractive product cards, and production deployment.",
        ],
    ):
        story.append(item)
    for item in project_block(
        "IdeaVault - Idea Management Web App",
        "React/Next.js, Node.js, MongoDB | Live: idea-vault-client-sigma.vercel.app",
        [
            "Developed a secure idea management platform where users can save, organize, and manage project ideas.",
            "Client: github.com/GalibDev/idea-vault-client | Server: github.com/GalibDev/idea-vault-server",
        ],
    ):
        story.append(item)
    for item in project_block(
        "Gatrix Robotics Website",
        "Modern responsive website | Live: gatrix.xyz",
        [
            "Created a polished robotics/business website with clear content sections, responsive layout, and deployment-ready UI.",
        ],
    ):
        story.append(item)

    add_section(story, "Experience", styles)
    story.extend(
        project_block(
            "Freelance / Local Client Web Projects",
            "Full Stack / Frontend Developer | 2024-Present",
            [
                "Designed and delivered 2-3 small websites for local clients, including responsive layouts and deployment support.",
                "Worked directly with clients to understand requirements, refine UI, and publish finished websites.",
            ],
        )
    )

    add_section(story, "Education", styles)
    education = [
        ("BSc in CSE", "Shyamoli Engineering College, Affiliated with University of Dhaka | 2024-Present"),
        ("HSC - Science", "Rangpur Model College, Rangpur | 2023-2024 | GPA 5.00"),
        ("SSC - Science", "Shishu Niketon High School | 2021-2022 | GPA 5.00"),
    ]
    for degree, meta in education:
        story.append(Paragraph(f"<b>{degree}</b> <font color='#666666'>| {meta}</font>", styles["role"]))

    add_section(story, "Additional", styles)
    story.append(
        Paragraph(
            "Languages: Bangla, English | Interests: full stack development, SaaS products, e-commerce, and robotics websites.",
            styles["body"],
        )
    )

    doc.build(story)
    PUBLIC_PDF_PATH.write_bytes(PDF_PATH.read_bytes())
    print(PDF_PATH)
    print(PUBLIC_PDF_PATH)


if __name__ == "__main__":
    make_resume()
    make_pdf_resume()
